import streamlit as st
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
import os
import tempfile
from io import BytesIO
from docx import Document
import re

st.set_page_config(page_title="PDF/Gambar ke Word Converter", layout="centered")

st.title("📄 PDF/Gambar ke Word Converter (Sederhana & Lanjutan)")

mode = st.radio("Pilih mode konversi:", ["Sederhana (otomatis)", "Lanjutan (pilih teks)"])

uploaded_file = st.file_uploader("Unggah file PDF atau gambar (JPG, PNG):", type=["pdf", "jpg", "jpeg", "png"])

def sanitize_text(text):
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b-\x1f\x7f]', '', text)
    return text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')

def extract_text_from_image(image):
    text = pytesseract.image_to_string(image)
    paragraphs = [para for para in text.split("\n") if para.strip()]
    return paragraphs

def display_docx_content(doc_buffer):
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    try:
        doc = Document(doc_buffer)
        st.markdown("### 📄 Pratinjau Isi Dokumen Word")
        for para in doc.paragraphs:
            if para.text.strip():
                st.write(para.text)
    except PackageNotFoundError:
        st.error("Gagal membuka dokumen Word untuk pratinjau.")

def checkbox_group(label, options, default=[], key_prefix=""):
    st.markdown(f"**{label}**")
    col1, col2 = st.columns([1, 1])
    with col1:
        all_state = st.button("✅ Pilih Semua", key=key_prefix + "_all")
    with col2:
        clear_state = st.button("❌ Kosongkan Semua", key=key_prefix + "_none")

    result = []
    for i, option in enumerate(options):
        option_key = f"{key_prefix}_{i}"
        checked = option in default
        if all_state:
            checked = True
        elif clear_state:
            checked = False
        if st.checkbox(option, key=option_key, value=checked):
            result.append(option)
    return result

if uploaded_file:
    file_name = uploaded_file.name.lower()

    if mode == "Sederhana (otomatis)":
        use_ocr = st.checkbox("Aktifkan OCR (untuk file hasil scan/foto/gambar)", value=False)

        if st.button("🔁 Convert to Word"):
            doc = Document()
            try:
                if file_name.endswith((".jpg", ".jpeg", ".png")):
                    image = Image.open(uploaded_file)
                    text = pytesseract.image_to_string(image)
                    clean_text = sanitize_text(text)
                    doc.add_paragraph(clean_text)

                elif file_name.endswith(".pdf"):
                    uploaded_file.seek(0)
                    if use_ocr:
                        import pdf2image
                        images = pdf2image.convert_from_bytes(uploaded_file.read())
                        for img in images:
                            text = pytesseract.image_to_string(img)
                            clean_text = sanitize_text(text)
                            doc.add_paragraph(clean_text)
                    else:
                        with pdfplumber.open(uploaded_file) as pdf:
                            for page in pdf.pages:
                                text = page.extract_text()
                                if text:
                                    clean_text = sanitize_text(text)
                                    doc.add_paragraph(clean_text)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                output_name = os.path.splitext(uploaded_file.name)[0] + " (konversi).docx"
                st.download_button("📥 Download Word", data=buffer, file_name=output_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            except Exception as e:
                st.error(f"Terjadi kesalahan saat konversi: {e}")

    elif mode == "Lanjutan (pilih teks)":
        use_ocr = st.checkbox("Aktifkan OCR untuk gambar (misal screenshot kode)", value=True)

        if file_name.endswith(".pdf"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_pdf_path = tmp_file.name

            with pdfplumber.open(tmp_pdf_path) as pdf:
                total_pages = len(pdf.pages)
                st.info(f"📄 File memiliki {total_pages} halaman.")

                selected_pages = st.multiselect("Pilih halaman yang ingin dikonversi:", list(range(1, total_pages + 1)), default=list(range(1, total_pages + 1)))

                if selected_pages:
                    all_extracted_text = []
                    with fitz.open(tmp_pdf_path) as doc:
                        for page_num in selected_pages:
                            page = pdf.pages[page_num - 1]
                            text = page.extract_text() or ""
                            img_text = ""
                            images = page.images
                            if use_ocr and (not text.strip() or images):
                                pix = doc.load_page(page_num - 1).get_pixmap()
                                image = Image.open(BytesIO(pix.tobytes()))
                                img_text = pytesseract.image_to_string(image)
                            combined_text = text + "\n" + img_text if img_text else text
                            all_extracted_text.append((page_num, combined_text.strip()))

                    st.markdown("### 🔍 Pratinjau & Pilih Teks")
                    selected_text = []
                    for page_num, text in all_extracted_text:
                        with st.expander(f"Halaman {page_num}"):
                            paragraphs = text.split("\n")
                            options = [para for para in paragraphs if para.strip()]
                            chosen = checkbox_group("Teks dari halaman ini:", options, key_prefix=f"p{page_num}")
                            selected_text.extend(chosen)

                    if selected_text:
                        doc = Document()
                        for para in selected_text:
                            doc.add_paragraph(para)
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        display_docx_content(buffer)
                        output_name = os.path.splitext(uploaded_file.name)[0] + " (konversi).docx"
                        st.download_button("⬇️ Unduh Hasil Word", data=buffer, file_name=output_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.warning("Pilih minimal satu halaman.")

        elif file_name.endswith((".jpg", ".jpeg", ".png")):
            image = Image.open(uploaded_file)
            st.image(image, caption="Gambar yang diunggah", use_column_width=True)
            paragraphs = extract_text_from_image(image)
            st.markdown("### 🔍 Pratinjau Teks Hasil OCR")
            selected_text = checkbox_group("Pilih teks yang ingin dikonversi:", paragraphs, key_prefix="ocr")
            if selected_text:
                doc = Document()
                for para in selected_text:
                    doc.add_paragraph(para)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                display_docx_content(buffer)
                output_name = os.path.splitext(uploaded_file.name)[0] + " (konversi).docx"
                st.download_button("⬇️ Unduh Hasil Word", data=buffer, file_name=output_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
